"""Unit tests for document_parser.py — parse PDF, DOCX, TXT, MD files."""
import os
import tempfile

import pytest

from app.services.document_parser import SUPPORTED_EXTENSIONS, parse_file


@pytest.fixture
def tmpdir():
    d = tempfile.mkdtemp()
    yield d
    import shutil
    shutil.rmtree(d, ignore_errors=True)


def _write(directory: str, name: str, content: str, encoding: str = "utf-8") -> str:
    path = os.path.join(directory, name)
    with open(path, "w", encoding=encoding) as f:
        f.write(content)
    return path


class TestParseTxt:
    def test_utf8(self, tmpdir):
        path = _write(tmpdir, "test.txt", "Hello 你好\n第二行")
        text, meta = parse_file(path)
        assert "Hello 你好" in text
        assert meta["file_type"] == "txt"
        assert meta["file_name"] == "test.txt"

    def test_latin1_fallback(self, tmpdir):
        path = os.path.join(tmpdir, "latin1.txt")
        with open(path, "w", encoding="latin-1") as f:
            f.write("café")
        text, _ = parse_file(path)
        assert "café" in text


class TestParseMarkdown:
    def test_markdown(self, tmpdir):
        path = _write(tmpdir, "readme.md", "# Title\n\nBody text")
        text, meta = parse_file(path)
        assert "# Title" in text
        assert meta["file_type"] == "md"


class TestParseDocx:
    def test_docx(self, tmpdir):
        pytest.importorskip("docx")
        from docx import Document

        path = os.path.join(tmpdir, "test.docx")
        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("第二段内容")
        doc.save(path)

        text, meta = parse_file(path)
        assert "第一段" in text
        assert "第二段内容" in text
        assert meta["file_type"] == "docx"


class TestUnsupported:
    def test_unsupported_extension(self, tmpdir):
        path = _write(tmpdir, "image.png", "fake")
        with pytest.raises(ValueError):
            parse_file(path)


class TestSupportedExtensions:
    def test_common_extensions(self):
        assert ".pdf" in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".txt" in SUPPORTED_EXTENSIONS
        assert ".md" in SUPPORTED_EXTENSIONS
        assert ".markdown" in SUPPORTED_EXTENSIONS
