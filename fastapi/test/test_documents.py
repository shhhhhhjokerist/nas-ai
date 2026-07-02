"""Unit tests for document parsing and chunking."""

import os
import tempfile
import unittest

from app.services.document_parser import SUPPORTED_EXTENSIONS, parse_file
from app.services.chunker import chunk_text


class TestDocumentParser(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _write(self, name: str, content: str) -> str:
        path = os.path.join(self.tmpdir, name)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    # ── txt ──

    def test_parse_txt_utf8(self):
        path = self._write("test.txt", "Hello 你好\n第二行")
        text, meta = parse_file(path)
        self.assertIn("Hello 你好", text)
        self.assertEqual(meta["file_type"], "txt")
        self.assertEqual(meta["file_name"], "test.txt")

    def test_parse_markdown(self):
        path = self._write("readme.md", "# Title\n\nBody text")
        text, meta = parse_file(path)
        self.assertIn("# Title", text)
        self.assertEqual(meta["file_type"], "md")

    def test_parse_txt_latin1_fallback(self):
        path = os.path.join(self.tmpdir, "latin1.txt")
        with open(path, "w", encoding="latin-1") as f:
            f.write("café")
        text, _ = parse_file(path)
        self.assertIn("café", text)

    def test_unsupported_extension(self):
        path = self._write("image.png", "fake")
        with self.assertRaises(ValueError):
            parse_file(path)

    def test_supported_extensions_set(self):
        self.assertIn(".pdf", SUPPORTED_EXTENSIONS)
        self.assertIn(".docx", SUPPORTED_EXTENSIONS)
        self.assertIn(".txt", SUPPORTED_EXTENSIONS)
        self.assertIn(".md", SUPPORTED_EXTENSIONS)
        self.assertIn(".markdown", SUPPORTED_EXTENSIONS)

    # ── docx (unit-like — requires python-docx) ──

    def test_parse_docx(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed")

        path = os.path.join(self.tmpdir, "test.docx")
        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("第二段内容")
        doc.save(path)

        text, meta = parse_file(path)
        self.assertIn("第一段", text)
        self.assertIn("第二段内容", text)
        self.assertEqual(meta["file_type"], "docx")


class TestChunker(unittest.TestCase):
    def test_empty_text(self):
        self.assertEqual(chunk_text(""), [])

    def test_short_text_single_chunk(self):
        chunks = chunk_text("短文本", chunk_size=500, chunk_overlap=50)
        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0], "短文本")

    def test_long_text_multiple_chunks(self):
        # Generate text longer than chunk_size
        text = "测试文本。" * 500  # ~2500 chars
        chunks = chunk_text(text, chunk_size=500, chunk_overlap=50)
        self.assertGreater(len(chunks), 1)
        for c in chunks:
            self.assertLessEqual(len(c), 500 + 100)  # allow some slack

    def test_chunk_overlap(self):
        text = (("段落%d。" % i) * 3 + "\n\n") * 20
        chunks = chunk_text(text, chunk_size=200, chunk_overlap=30)
        if len(chunks) >= 2:
            # Last bit of chunk 0 should appear at start of chunk 1
            overlap_part = chunks[0][-20:]
            self.assertIn(overlap_part[:5], chunks[1])


if __name__ == "__main__":
    unittest.main()
